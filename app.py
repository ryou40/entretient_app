
import os
import streamlit as st
import docx
from docx import Document
import openai
import json
import re
from io import BytesIO
from datetime import date
import copy
import unicodedata

# --- CONFIGURATION & CLE ---
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

STRUCTURE_JSON = {
    "disponibilite": "non précisé",
    "base": "non précisé",
    "mobilite": "non précisé",
    "permis": "non précisé",
    "vehicule": "non précisé",
    "remuneration_souhaitee": "non précisé",
    "poste_souhaite": "non précisé",
    "secteurs_souhaites": "non précisé",
    "entreprises_souhaitees": "non précisé",
    "evolution_souhaitee": "non précisé",
    "criteres_choix": [],
    "logiciels": {},
    "langues": {
        "anglais_oral": "non précisé",
        "autres_langues": "non précisé"
    },
    "references": "non précisé",
    "loisirs": "non précisé",
    "projets_experiences": [],
    "infos_complementaires": {
        "changement_statut": "non précisé",
        "famille": "non précisé",
        "process_avance_qualif": "non précisé"
    }
}

# --- Fonctions de traitement ---

def demande_segment_avec_coupure(segment: str, numero_segment: int) -> str:
    """
    Envoie un segment de texte à GPT pour trouver un point de coupure naturelle.
    Retourne le texte original jusqu'au point de coupure, suivi de la balise <<<COUPURE>>>.
    """
    prompt = (
        "Tu vas m’aider à découper une transcription d’entretien en morceaux cohérents.\n"
        "Je t’envoie un extrait. Réécris exactement le texte d’origine depuis le début jusqu’à un endroit naturel pour couper "
        "(juste avant un changement de sujet, une question ou une pause logique).\n"
        "⚠️ Ne dépasse pas cet endroit. Ajoute exactement cette balise à la fin : <<<COUPURE>>>.\n"
        "N’invente rien, ne reformule pas et ne commente pas. Copie-colle uniquement le texte jusqu’à la coupure demandée.\n\n"
        f"{segment}"
    )
    # Appel à l’API OpenAI (modèle GPT-4) pour obtenir la coupure
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    resultat = response.choices[0].message.content.strip()
    return resultat  # Contient le segment d’origine jusqu’au point de coupure, terminé par <<<COUPURE>>>

def decoupe_transcription(texte: str, min_size: int = 2000, max_size: int = 4000) -> list:
    """
    Découpe un texte de transcription en blocs cohérents en utilisant GPT pour trouver les points de coupure.
    Renvoie une liste de blocs de texte (chaque bloc étant une portion cohérente de la transcription).
    """
    blocs = []
    cursor = 0
    total_length = len(texte)
    bloc_num = 1

    while cursor < total_length:
        reste = total_length - cursor
        # Si le reste du texte est plus petit que la taille max, on le prend d'un coup
        if reste < max_size:
            bloc_final = texte[cursor:]
            blocs.append(bloc_final)
            # Dernier bloc ajouté, fin de la boucle
            break

        # Définir une fenêtre [min_size, max_size] pour envoyer un segment à GPT
        window_start = cursor + min_size
        window_end = min(cursor + max_size, total_length)
        segment = texte[window_start:window_end]

        try:
            # Demander à GPT le segment avec marque de coupure
            reponse = demande_segment_avec_coupure(segment, bloc_num)
            coupure_index = reponse.find("<<<COUPURE>>>")
            if coupure_index == -1:
                # Si la balise n’a pas été trouvée dans la réponse GPT, on considère une erreur
                raise ValueError("Balise <<<COUPURE>>> non trouvée dans la réponse GPT.")
        except Exception as e:
            # En cas d’erreur (API non disponible, balise absente, etc.), on coupe intelligemment à la fin du segment
            print(f"⚠️ Erreur lors de la demande de coupure (bloc {bloc_num}) : {e}")

            # Ajustement de la zone de recherche (max 300 derniers caractères)
            zone_recherche = segment[-min(len(segment), 300):]

            # Recherche d'une fin de phrase suivie d'une majuscule (ou début logique)
            candidates = list(re.finditer(r'([.!?…])\s+[A-ZÀÂÉÈÊËÎÏÔÖÙÛÜÇ]', zone_recherche))
            if candidates:
                # Prendre la dernière ponctuation suivie d’un début de phrase
                cut_offset = candidates[-1].end()
                coupure_index = len(segment) - len(zone_recherche) + cut_offset
                print(f"🔪 Coupure par ponctuation trouvée à l'index {coupure_index}")
            else:
                # Sinon, couper à l'espace le plus proche autour de 1000
                last_space = segment.rfind(' ', 1000)
                coupure_index = last_space if last_space != -1 else len(segment)
                print(f"🔪 Coupure fallback à l'espace à l'index {coupure_index}")

            reponse = segment[:coupure_index]


        # Calcul du point de coupure absolu dans le texte original
        cut_point = window_start + coupure_index
        # Extraire le bloc du texte original depuis le curseur actuel jusqu'au point de coupure
        bloc = texte[cursor:cut_point]
        blocs.append(bloc)

        # Mettre à jour le curseur pour le prochain bloc
        cursor = cut_point
        bloc_num += 1

    return blocs

def extraire_noms_interviewers(bloc_texte: str) -> list:
    prompt = f"""
    Voici un extrait de transcription d'entretien d'embauche.
    
    Les interventions sont identifiables par leurs *speaker labels*.
    
    Ton objectif est de :
    1. Identifier la personne de l’entreprise (RH ou recruteur) qui mène l’entretien.
    2. Identifier le ou la candidat(e) interrogé(e).
    
    Retourne une **liste JSON de deux éléments**, dans cet ordre :
    1. Le **trigramme** du recruteur : 1ère lettre du prénom + 2 premières lettres du nom (en MAJUSCULES).
    2. Le **speaker label exact** du candidat, tel qu’il apparaît dans la transcription.
    
    ⚠️ Ne retourne **rien d’autre** que cette liste JSON.
    
    Exemple de réponse attendue :
    ["JMA", "Nina Dubois"]

    
    Texte :
    \"\"\"{bloc_texte}\"\"\"
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            temperature=0,
            messages=[{"role": "user", "content": prompt}]
        )
        content = response.choices[0].message.content.strip()

        # Capturer uniquement la première liste JSON trouvée
        match = re.search(r'\[\s*".+?"\s*,\s*".+?"\s*\]', content, re.DOTALL)
        if match:
            raw_json = match.group(0)
            resultat = json.loads(raw_json)

            # Vérification + nettoyage de chaque élément
            if (
                isinstance(resultat, list)
                and len(resultat) == 2
                and all(isinstance(x, str) for x in resultat)
            ):
                return [x.strip().strip(".").strip() for x in resultat]
            else:
                print("⚠️ Format JSON incorrect ou incomplet :", resultat)
        else:
            print("⚠️ Aucune liste JSON trouvée dans la réponse GPT.")

    except Exception as e:
        print(f"❌ Erreur lors de l'extraction des interviewers : {e}")
    
    return []



def extraire_infos_depuis_texte(texte_transcrit: str) -> dict:
    """
    Utilise GPT pour extraire les informations clés d'une portion de transcription d'entretien.
    Retourne un dictionnaire avec les champs d'information requis.
    """
    # Préparation du prompt pour GPT
    prompt = f"""
    Voici la transcription d'un entretien d'embauche :
    Analyse la transcription et retourne exclusivement un objet JSON avec les informations ci-dessous.
    Si une information n'est pas mentionnée dans la transcription, indique "non précisé".
    Ne commente pas, ne reformule pas. Retourne uniquement l'objet JSON, sans texte autour.

    Si le candidat parle d’un projet ou d’une expérience professionnelle, rédige une synthèse courte et complète de ce qu’il en dit : poste, contexte, type de projet, montant, enjeux, responsabilités, résultats, etc., selon ce qui est mentionné. Liste-les dans `projets_experiences`.

    Champs attendus :
    {json.dumps(STRUCTURE_JSON, indent=2, ensure_ascii=False)}

    La transcription :
    \"\"\"{texte_transcrit}\"\"\"
    """
    try:
        # Appel API OpenAI (GPT-4) pour obtenir l'objet JSON demandé
        response = client.chat.completions.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {"role": "system", "content": "Tu es un assistant RH qui structure des données d'entretien en JSON."},
                {"role": "user", "content": prompt}
            ]
        )
        content = response.choices[0].message.content.strip()
    except Exception as e:
        # En cas d'échec de l'appel API, on retourne un dict vide
        print(f"❌ Erreur lors de l'appel à l'API OpenAI : {e}")
        st.write("🔧 Debug GPT brut :")
        st.json(response.model_dump())

        return {}

    # Extraction du texte JSON de la réponse (pour ignorer d'éventuels commentaires ou textes parasites)
    match = re.search(r'\{.*\}', content, re.DOTALL)
    if match:
        json_text = match.group(0)
        try:
            return json.loads(json_text)
        except json.JSONDecodeError:
            print("❌ Erreur : JSON invalide obtenu.")
            return {}
    else:
        print("❌ Erreur : aucun objet JSON trouvé dans la réponse GPT.")
        return {}

def fusionner_infos(global_infos: dict, nouvelles_infos: dict) -> None:
    """
    Fusionne un dictionnaire d'infos extraites (`nouvelles_infos`) avec le dictionnaire `global_infos` (cumulatif).
    Règles de fusion :
      - Si le champ global est "non précisé" ou vide, on le remplace par la nouvelle valeur.
      - Si le champ global a une valeur différente et que la nouvelle valeur est également renseignée (≠ "non précisé"),
        on transforme le champ en liste de valeurs (sans doublons) pour conserver toutes les informations.
      - Pour les sous-dictionnaires et listes, la fusion est appliquée récursivement (union de clés, listes fusionnées).
    """
    for cle, valeur in nouvelles_infos.items():
        if isinstance(valeur, dict):
            # Fusionner récursivement les dictionnaires imbriqués
            if cle not in global_infos or not isinstance(global_infos.get(cle), dict):
                global_infos[cle] = {}
            fusionner_infos(global_infos[cle], valeur)
        elif isinstance(valeur, list):
            # Fusionner les listes en ajoutant les nouveaux éléments non duplicatifs
            if cle not in global_infos or not isinstance(global_infos.get(cle), list):
                global_infos[cle] = []
            for item in valeur:
                if item not in global_infos[cle]:
                    global_infos[cle].append(item)
        else:
            # Fusionner les champs scalaires
            if cle not in global_infos or global_infos[cle] == "non précisé":
                global_infos[cle] = valeur
            elif global_infos[cle] != valeur and valeur != "non précisé":
                # Si conflit et aucune des deux valeurs n'est "non précisé", stocker sous forme de liste
                if not isinstance(global_infos[cle], list):
                    global_infos[cle] = [global_infos[cle]]
                if valeur not in global_infos[cle]:
                    global_infos[cle].append(valeur)

def ajouter_contenu(doc, titre, contenu, niveau=1):
    """Ajoute récursivement une section dans le document Word selon le type du contenu."""
    if isinstance(contenu, dict):
        doc.add_heading(titre.replace("_", " ").capitalize(), level=niveau)
        for cle, val in contenu.items():
            ajouter_contenu(doc, cle, val, niveau + 1)

    elif isinstance(contenu, list):
        doc.add_heading(titre.replace("_", " ").capitalize(), level=niveau)
        for i, item in enumerate(contenu, 1):
            sous_titre = f"{titre[:-1]} {i}" if isinstance(item, dict) else f"{item}"
            ajouter_contenu(doc, sous_titre, item, niveau + 1)

    else:
        texte = f"{titre.replace('_', ' ').capitalize()} : {contenu}"
        doc.add_paragraph(texte, style="List Bullet")


def generer_docx(infos: dict, interviewers: list) -> BytesIO:
    doc = Document()
    doc.add_heading("Compte-rendu d'entretien", 0)

    # Sous-titre avec les interviewers
    if interviewers:
        doc.add_paragraph("Intervenants : " + ", ".join(interviewers))

    doc.add_paragraph(" ")

    # 🔁 Nouvelle logique générique pour tout afficher proprement
    for cle, valeur in infos.items():
        ajouter_contenu(doc, cle, valeur, niveau=2)

    # Enregistrement en mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

import unicodedata

def slugify_nom(nom: str) -> str:
    """
    Nettoie un nom pour usage en nom de fichier :
    - enlève les accents
    - supprime les caractères spéciaux
    - remplace les espaces par des underscores
    - garde uniquement lettres, chiffres, et _
    """
    # Normalise les caractères accentués
    nom = unicodedata.normalize('NFKD', nom)
    nom = nom.encode('ascii', 'ignore').decode('utf-8')

    # Remplace les espaces par _
    nom = nom.replace(" ", "_")

    # Supprime tout sauf lettres, chiffres, _
    nom = re.sub(r'[^A-Za-z0-9_]', '', nom)

    return nom

# --- INTERFACE STREAMLIT ---
st.title("Analyse de transcription d'entretien")
st.write("Importez un fichier Word contenant la transcription de l’entretien, puis cliquez sur **Traiter** pour lancer l’analyse.")
debug = st.checkbox("🔧 Mode développeur")

fichier_docx = st.file_uploader("Transcription d'entretien (.docx)", type=['docx'])

if fichier_docx is not None:
    if "processing" not in st.session_state:
        st.session_state.processing = False

    if st.button("Traiter", disabled=st.session_state.processing):
        if not st.session_state.processing:
            st.session_state.processing = True

            with st.spinner("⏳ Traitement en cours… Cela peut prendre quelques minutes."):
                doc = docx.Document(fichier_docx)
                texte_complet = "\n".join([para.text for para in doc.paragraphs])
                if debug:
                    st.write("📄 Longueur texte :", len(texte_complet))
                    st.write("🔍 Début texte :", texte_complet[:500])

                blocs = decoupe_transcription(texte_complet)
                if debug:
                    st.write("📦 Nombre de blocs générés :", len(blocs))
                    st.write("🧩 Bloc 1 :", blocs[0][:500] if blocs else "Aucun")

                noms_interviewers = extraire_noms_interviewers(blocs[0])
                st.write("👥 Interviewers détectés :", noms_interviewers)

                global_infos = copy.deepcopy(STRUCTURE_JSON)

                for i, bloc in enumerate(blocs):
                    if debug:
                        st.write(f"--- 🔄 Bloc {i+1}/{len(blocs)} ---")
                        st.write(bloc[:300])

                    infos_extraites = extraire_infos_depuis_texte(bloc)
                    if debug:
                        st.write("🧠 Infos extraites :", infos_extraites)

                    fusionner_infos(global_infos, infos_extraites)

            if debug:
                st.subheader("🧾 Résumé structuré (debug)")
                st.json(global_infos)

            buffer = generer_docx(global_infos, noms_interviewers)
            aujourd_hui = date.today().strftime("%Y-%m-%d")

            trigramme_rh = "RH"
            nom_candidat = "candidat"

            if isinstance(noms_interviewers, list):
                if len(noms_interviewers) > 0 and noms_interviewers[0].strip():
                    trigramme_rh = noms_interviewers[0].strip().upper()
                if len(noms_interviewers) > 1 and noms_interviewers[1].strip():
                    nom_candidat = slugify_nom(noms_interviewers[1].strip())

            nom_fichier = f"e1__{nom_candidat}-{trigramme_rh}__{aujourd_hui}.docx"

            st.success("✅ Analyse terminée. Vous pouvez télécharger le compte-rendu ci-dessous.")
            st.download_button(
                label="📥 Télécharger le document Word",
                data=buffer,
                file_name=nom_fichier,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.session_state.processing = False

    else:
        st.info("Cliquez sur le bouton **Traiter** pour lancer l’analyse de la transcription.")
else:
    st.warning("Veuillez importer un fichier Word (.docx) pour commencer.")


# --- LES FONCTIONS (copiées telles quelles, déjà correctes) ---
# Ajoute ici toutes tes fonctions précédentes telles que :
# - demande_segment_avec_coupure
# - decoupe_transcription
# - extraire_noms_interviewers
# - extraire_infos_depuis_texte
# - fusionner_infos
# - ajouter_contenu
# - generer_docx
# - slugify_nom

# (Je ne les recopie pas ici pour éviter de surcharger, mais tu peux juste les coller en bas du fichier)




#import os
#import streamlit as st
#import docx                     # Lecture de documents Word
#from docx import Document
#import openai                   # Appels API OpenAI GPT
#from dotenv import load_dotenv
#import json
#import re
#from io import BytesIO
#from datetime import date
#import copy
#
## Chargement de la clé API OpenAI depuis un fichier .env ou les variables d'environnement
#client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
#
#STRUCTURE_JSON = {
#    "disponibilite": "non précisé",
#    "base": "non précisé",
#    "mobilite": "non précisé",
#    "permis": "non précisé",
#    "vehicule": "non précisé",
#    "remuneration_souhaitee": "non précisé",
#    "poste_souhaite": "non précisé",
#    "secteurs_souhaites": "non précisé",
#    "entreprises_souhaitees": "non précisé",
#    "evolution_souhaitee": "non précisé",
#    "criteres_choix": [],
#    "logiciels": {},
#    "langues": {
#        "anglais_oral": "non précisé",
#        "autres_langues": "non précisé"
#    },
#    "references": "non précisé",
#    "loisirs": "non précisé",
#    "projets_experiences": [],
#    "infos_complementaires": {
#        "changement_statut": "non précisé",
#        "famille": "non précisé",
#        "process_avance_qualif": "non précisé"
#    }
#}
#
## --- Fonctions de traitement ---
#
#def demande_segment_avec_coupure(segment: str, numero_segment: int) -> str:
#    """
#    Envoie un segment de texte à GPT pour trouver un point de coupure naturelle.
#    Retourne le texte original jusqu'au point de coupure, suivi de la balise <<<COUPURE>>>.
#    """
#    prompt = (
#        "Tu vas m’aider à découper une transcription d’entretien en morceaux cohérents.\n"
#        "Je t’envoie un extrait. Réécris exactement le texte d’origine depuis le début jusqu’à un endroit naturel pour couper "
#        "(juste avant un changement de sujet, une question ou une pause logique).\n"
#        "⚠️ Ne dépasse pas cet endroit. Ajoute exactement cette balise à la fin : <<<COUPURE>>>.\n"
#        "N’invente rien, ne reformule pas et ne commente pas. Copie-colle uniquement le texte jusqu’à la coupure demandée.\n\n"
#        f"{segment}"
#    )
#    # Appel à l’API OpenAI (modèle GPT-4) pour obtenir la coupure
#    response = client.chat.completions.create(
#        model="gpt-4",
#        messages=[{"role": "user", "content": prompt}],
#        temperature=0
#    )
#    resultat = response.choices[0].message.content.strip()
#    return resultat  # Contient le segment d’origine jusqu’au point de coupure, terminé par <<<COUPURE>>>
#
#def decoupe_transcription(texte: str, min_size: int = 2000, max_size: int = 4000) -> list:
#    """
#    Découpe un texte de transcription en blocs cohérents en utilisant GPT pour trouver les points de coupure.
#    Renvoie une liste de blocs de texte (chaque bloc étant une portion cohérente de la transcription).
#    """
#    blocs = []
#    cursor = 0
#    total_length = len(texte)
#    bloc_num = 1
#
#    while cursor < total_length:
#        reste = total_length - cursor
#        # Si le reste du texte est plus petit que la taille max, on le prend d'un coup
#        if reste < max_size:
#            bloc_final = texte[cursor:]
#            blocs.append(bloc_final)
#            # Dernier bloc ajouté, fin de la boucle
#            break
#
#        # Définir une fenêtre [min_size, max_size] pour envoyer un segment à GPT
#        window_start = cursor + min_size
#        window_end = min(cursor + max_size, total_length)
#        segment = texte[window_start:window_end]
#
#        try:
#            # Demander à GPT le segment avec marque de coupure
#            reponse = demande_segment_avec_coupure(segment, bloc_num)
#            coupure_index = reponse.find("<<<COUPURE>>>")
#            if coupure_index == -1:
#                # Si la balise n’a pas été trouvée dans la réponse GPT, on considère une erreur
#                raise ValueError("Balise <<<COUPURE>>> non trouvée dans la réponse GPT.")
#        except Exception as e:
#            # En cas d’erreur (API non disponible, balise absente, etc.), on coupe intelligemment à la fin du segment
#            print(f"⚠️ Erreur lors de la demande de coupure (bloc {bloc_num}) : {e}")
#
#            # Ajustement de la zone de recherche (max 300 derniers caractères)
#            zone_recherche = segment[-min(len(segment), 300):]
#
#            # Recherche d'une fin de phrase suivie d'une majuscule (ou début logique)
#            candidates = list(re.finditer(r'([.!?…])\s+[A-ZÀÂÉÈÊËÎÏÔÖÙÛÜÇ]', zone_recherche))
#            if candidates:
#                # Prendre la dernière ponctuation suivie d’un début de phrase
#                cut_offset = candidates[-1].end()
#                coupure_index = len(segment) - len(zone_recherche) + cut_offset
#                print(f"🔪 Coupure par ponctuation trouvée à l'index {coupure_index}")
#            else:
#                # Sinon, couper à l'espace le plus proche autour de 1000
#                last_space = segment.rfind(' ', 1000)
#                coupure_index = last_space if last_space != -1 else len(segment)
#                print(f"🔪 Coupure fallback à l'espace à l'index {coupure_index}")
#
#            reponse = segment[:coupure_index]
#
#
#        # Calcul du point de coupure absolu dans le texte original
#        cut_point = window_start + coupure_index
#        # Extraire le bloc du texte original depuis le curseur actuel jusqu'au point de coupure
#        bloc = texte[cursor:cut_point]
#        blocs.append(bloc)
#
#        # Mettre à jour le curseur pour le prochain bloc
#        cursor = cut_point
#        bloc_num += 1
#
#    return blocs
#
#def extraire_noms_interviewers(bloc_texte: str) -> list:
#    prompt = f"""
#    Voici un extrait de transcription d'entretien d'embauche.
#    
#    Les interventions sont identifiables par leurs *speaker labels*.
#    
#    Ton objectif est de :
#    1. Identifier la personne de l’entreprise (RH ou recruteur) qui mène l’entretien.
#    2. Identifier le ou la candidat(e) interrogé(e).
#    
#    Retourne une **liste JSON de deux éléments**, dans cet ordre :
#    1. Le **trigramme** du recruteur : 1ère lettre du prénom + 2 premières lettres du nom (en MAJUSCULES).
#    2. Le **speaker label exact** du candidat, tel qu’il apparaît dans la transcription.
#    
#    ⚠️ Ne retourne **rien d’autre** que cette liste JSON.
#    
#    Exemple de réponse attendue :
#    ["JMA", "Nina Dubois"]
#
#    
#    Texte :
#    \"\"\"{bloc_texte}\"\"\"
#    """
#    try:
#        response = client.chat.completions.create(
#            model="gpt-4",
#            temperature=0,
#            messages=[{"role": "user", "content": prompt}]
#        )
#        content = response.choices[0].message.content.strip()
#
#        # Capturer uniquement la première liste JSON trouvée
#        match = re.search(r'\[\s*".+?"\s*,\s*".+?"\s*\]', content, re.DOTALL)
#        if match:
#            raw_json = match.group(0)
#            resultat = json.loads(raw_json)
#
#            # Vérification + nettoyage de chaque élément
#            if (
#                isinstance(resultat, list)
#                and len(resultat) == 2
#                and all(isinstance(x, str) for x in resultat)
#            ):
#                return [x.strip().strip(".").strip() for x in resultat]
#            else:
#                print("⚠️ Format JSON incorrect ou incomplet :", resultat)
#        else:
#            print("⚠️ Aucune liste JSON trouvée dans la réponse GPT.")
#
#    except Exception as e:
#        print(f"❌ Erreur lors de l'extraction des interviewers : {e}")
#    
#    return []
#
#
#
#def extraire_infos_depuis_texte(texte_transcrit: str) -> dict:
#    """
#    Utilise GPT pour extraire les informations clés d'une portion de transcription d'entretien.
#    Retourne un dictionnaire avec les champs d'information requis.
#    """
#    # Préparation du prompt pour GPT
#    prompt = f"""
#    Voici la transcription d'un entretien d'embauche :
#    Analyse la transcription et retourne exclusivement un objet JSON avec les informations ci-dessous.
#    Si une information n'est pas mentionnée dans la transcription, indique "non précisé".
#    Ne commente pas, ne reformule pas. Retourne uniquement l'objet JSON, sans texte autour.
#
#    Si le candidat parle d’un projet ou d’une expérience professionnelle, rédige une synthèse courte et complète de ce qu’il en dit : poste, contexte, type de projet, montant, enjeux, responsabilités, résultats, etc., selon ce qui est mentionné. Liste-les dans `projets_experiences`.
#
#    Champs attendus :
#    {json.dumps(STRUCTURE_JSON, indent=2, ensure_ascii=False)}
#
#    La transcription :
#    \"\"\"{texte_transcrit}\"\"\"
#    """
#    try:
#        # Appel API OpenAI (GPT-4) pour obtenir l'objet JSON demandé
#        response = client.chat.completions.create(
#            model="gpt-4",
#            temperature=0,
#            messages=[
#                {"role": "system", "content": "Tu es un assistant RH qui structure des données d'entretien en JSON."},
#                {"role": "user", "content": prompt}
#            ]
#        )
#        content = response.choices[0].message.content.strip()
#    except Exception as e:
#        # En cas d'échec de l'appel API, on retourne un dict vide
#        print(f"❌ Erreur lors de l'appel à l'API OpenAI : {e}")
#        return {}
#
#    # Extraction du texte JSON de la réponse (pour ignorer d'éventuels commentaires ou textes parasites)
#    match = re.search(r'\{.*\}', content, re.DOTALL)
#    if match:
#        json_text = match.group(0)
#        try:
#            return json.loads(json_text)
#        except json.JSONDecodeError:
#            print("❌ Erreur : JSON invalide obtenu.")
#            return {}
#    else:
#        print("❌ Erreur : aucun objet JSON trouvé dans la réponse GPT.")
#        return {}
#
#def fusionner_infos(global_infos: dict, nouvelles_infos: dict) -> None:
#    """
#    Fusionne un dictionnaire d'infos extraites (`nouvelles_infos`) avec le dictionnaire `global_infos` (cumulatif).
#    Règles de fusion :
#      - Si le champ global est "non précisé" ou vide, on le remplace par la nouvelle valeur.
#      - Si le champ global a une valeur différente et que la nouvelle valeur est également renseignée (≠ "non précisé"),
#        on transforme le champ en liste de valeurs (sans doublons) pour conserver toutes les informations.
#      - Pour les sous-dictionnaires et listes, la fusion est appliquée récursivement (union de clés, listes fusionnées).
#    """
#    for cle, valeur in nouvelles_infos.items():
#        if isinstance(valeur, dict):
#            # Fusionner récursivement les dictionnaires imbriqués
#            if cle not in global_infos or not isinstance(global_infos.get(cle), dict):
#                global_infos[cle] = {}
#            fusionner_infos(global_infos[cle], valeur)
#        elif isinstance(valeur, list):
#            # Fusionner les listes en ajoutant les nouveaux éléments non duplicatifs
#            if cle not in global_infos or not isinstance(global_infos.get(cle), list):
#                global_infos[cle] = []
#            for item in valeur:
#                if item not in global_infos[cle]:
#                    global_infos[cle].append(item)
#        else:
#            # Fusionner les champs scalaires
#            if cle not in global_infos or global_infos[cle] == "non précisé":
#                global_infos[cle] = valeur
#            elif global_infos[cle] != valeur and valeur != "non précisé":
#                # Si conflit et aucune des deux valeurs n'est "non précisé", stocker sous forme de liste
#                if not isinstance(global_infos[cle], list):
#                    global_infos[cle] = [global_infos[cle]]
#                if valeur not in global_infos[cle]:
#                    global_infos[cle].append(valeur)
#
#def ajouter_contenu(doc, titre, contenu, niveau=1):
#    """Ajoute récursivement une section dans le document Word selon le type du contenu."""
#    if isinstance(contenu, dict):
#        doc.add_heading(titre.replace("_", " ").capitalize(), level=niveau)
#        for cle, val in contenu.items():
#            ajouter_contenu(doc, cle, val, niveau + 1)
#
#    elif isinstance(contenu, list):
#        doc.add_heading(titre.replace("_", " ").capitalize(), level=niveau)
#        for i, item in enumerate(contenu, 1):
#            sous_titre = f"{titre[:-1]} {i}" if isinstance(item, dict) else f"{item}"
#            ajouter_contenu(doc, sous_titre, item, niveau + 1)
#
#    else:
#        texte = f"{titre.replace('_', ' ').capitalize()} : {contenu}"
#        doc.add_paragraph(texte, style="List Bullet")
#
#
#def generer_docx(infos: dict, interviewers: list) -> BytesIO:
#    doc = Document()
#    doc.add_heading("Compte-rendu d'entretien", 0)
#
#    # Sous-titre avec les interviewers
#    if interviewers:
#        doc.add_paragraph("Intervenants : " + ", ".join(interviewers))
#
#    doc.add_paragraph(" ")
#
#    # 🔁 Nouvelle logique générique pour tout afficher proprement
#    for cle, valeur in infos.items():
#        ajouter_contenu(doc, cle, valeur, niveau=2)
#
#    # Enregistrement en mémoire
#    buffer = BytesIO()
#    doc.save(buffer)
#    buffer.seek(0)
#    return buffer
#
#import unicodedata
#
#def slugify_nom(nom: str) -> str:
#    """
#    Nettoie un nom pour usage en nom de fichier :
#    - enlève les accents
#    - supprime les caractères spéciaux
#    - remplace les espaces par des underscores
#    - garde uniquement lettres, chiffres, et _
#    """
#    # Normalise les caractères accentués
#    nom = unicodedata.normalize('NFKD', nom)
#    nom = nom.encode('ascii', 'ignore').decode('utf-8')
#
#    # Remplace les espaces par _
#    nom = nom.replace(" ", "_")
#
#    # Supprime tout sauf lettres, chiffres, _
#    nom = re.sub(r'[^A-Za-z0-9_]', '', nom)
#
#    return nom
#
#
## --- Interface utilisateur Streamlit ---
#
#st.title("Analyse de transcription d'entretien")
#st.write("Importez un fichier Word contenant la transcription de l’entretien, puis cliquez sur **Traiter** pour lancer l’analyse.")
#
## Champ d'upload de fichier .docx
#fichier_docx = st.file_uploader("Transcription d'entretien (.docx)", type=['docx'])
#
## Bouton de traitement (n'apparaît que si un fichier est uploadé)
#if fichier_docx is not None:
## Initialiser le flag de traitement si non défini
#    if "processing" not in st.session_state:
#        st.session_state.processing = False
#
#    if st.button("Traiter", disabled=st.session_state.processing):
#        if not st.session_state.processing:
#            st.session_state.processing = True
#            with st.spinner("⏳ Traitement en cours… Cela peut prendre quelques minutes."):
#                # Lecture du fichier Word en mémoire
#                doc = docx.Document(fichier_docx)
#                texte_complet = "\n".join([para.text for para in doc.paragraphs])
#                blocs = decoupe_transcription(texte_complet)
#                # On extrait les noms des interviewers à partir du premier bloc
#                noms_interviewers = extraire_noms_interviewers(blocs[0])
#                st.write("👥 Interviewers détectés :", noms_interviewers)
#
#                global_infos = copy.deepcopy(STRUCTURE_JSON)
#
#                for bloc in blocs:
#                    infos_extraites = extraire_infos_depuis_texte(bloc)
#                    fusionner_infos(global_infos, infos_extraites)
#
#            buffer = generer_docx(global_infos, noms_interviewers)
#
#            aujourd_hui = date.today().strftime("%Y-%m-%d")
#            
#            # Par défaut
#            trigramme_rh = "RH"
#            nom_candidat = "candidat"
#
#            # Si GPT a renvoyé une liste, on essaye d’en extraire les 2 éléments
#            if isinstance(noms_interviewers, list):
#                if len(noms_interviewers) > 0 and noms_interviewers[0].strip():
#                    trigramme_rh = noms_interviewers[0].strip().upper()
#                if len(noms_interviewers) > 1 and noms_interviewers[1].strip():
#                    nom_candidat = slugify_nom(noms_interviewers[1].strip())
#
#            # Formatage final avec fallback
#            nom_fichier = f"e1__{nom_candidat}-{trigramme_rh}__{aujourd_hui}.docx"
#
#
#            st.success("✅ Analyse terminée. Vous pouvez télécharger le compte-rendu ci-dessous.")
#            st.download_button(
#                label="📥 Télécharger le document Word",
#                data=buffer,
#                file_name=nom_fichier,
#                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#            )
#            st.session_state.processing = False
#
#    else:
#        st.info("Cliquez sur le bouton **Traiter** pour lancer l’analyse de la transcription.")
#else:
#    st.warning("Veuillez importer un fichier Word (.docx) pour commencer.")